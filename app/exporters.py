"""将统一的讲稿与音频数据导出为文档、字幕或音频产物。

本模块不依赖界面状态；页面只负责选择目标路径并展示执行结果。
"""

import re
import shutil
import wave
from datetime import datetime
from pathlib import Path

from app import script_io


# WordprocessingML 不接受这些控制字符，写入单元格前统一过滤。
_CONTROL_CHARS_RE = re.compile('[' + ''.join(
    chr(c) for c in list(range(0, 9)) + [11, 12] + list(range(14, 32)) + [127]) + ']')


def _apply_font(run, *, size=11, bold=False, color=None):
    """为 Word 文本同时设置西文和东亚字体。"""
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(size)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def format_duration(seconds: float) -> str:
    """将秒数格式化为文档使用的 ``mm:ss``。"""
    total = int(round(max(seconds, 0)))
    minutes, secs = divmod(total, 60)
    return f'{minutes:02d}:{secs:02d}'


def format_srt_timestamp(seconds: float) -> str:
    """将秒数格式化为 SRT 时间戳 ``HH:MM:SS,mmm``。"""
    if seconds < 0:
        seconds = 0.0
    milliseconds = int(round(seconds * 1000))
    hours, milliseconds = divmod(milliseconds, 3600000)
    minutes, milliseconds = divmod(milliseconds, 60000)
    secs, milliseconds = divmod(milliseconds, 1000)
    return f'{hours:02d}:{minutes:02d}:{secs:02d},{milliseconds:03d}'


def group_by_page(notes_list, durations):
    """按相邻页码归组讲稿段，并保留输入顺序。"""
    grouped = []
    for index, note in enumerate(notes_list):
        duration = durations[index] if index < len(durations) else 0.0
        page = int(note.get('page', 0))
        if grouped and grouped[-1][0] == page:
            grouped[-1][1].append((note.get('text', ''), duration))
        else:
            grouped.append((page, [(note.get('text', ''), duration)]))
    return grouped


def page_summary(notes_list, durations):
    """按页汇总讲稿，返回 ``(页码, 分段列表, 时长, 字数)`` 序列。"""
    rows = []
    for page, segments in group_by_page(notes_list, durations):
        texts = [str(text) for text, _ in segments]
        duration = sum(item[1] for item in segments)
        words = sum(len(''.join(str(text).split())) for text in texts)
        rows.append((page, texts, duration, words))
    return rows


def build_meta_line(notes_list, durations, *, engine_name='', speaker='') -> str:
    """生成文档顶部使用的页数、段数、字数和语音信息摘要。"""
    rows = page_summary(notes_list, durations)
    total = sum(durations) if durations else 0.0
    words = sum(row[3] for row in rows)

    parts = [f'{len(rows)} 页', f'{len(notes_list)} 段', f'{words} 字']
    if total > 0:
        parts.append(f'预计 {format_duration(total)}')
    if engine_name:
        parts.append(f'引擎 {engine_name}')
    if speaker:
        parts.append(f'发音人 {speaker}')
    return ' / '.join(parts)


def build_markdown(notes_list, durations, *, title='演讲稿', mark='●',
                   engine_name='', speaker=''):
    """生成带全局摘要和逐页章节的只读 Markdown 讲稿。"""
    rows = page_summary(notes_list, durations)
    lines = [f'# {title}', '']
    lines.append('> ' + build_meta_line(notes_list, durations,
                                        engine_name=engine_name, speaker=speaker))
    lines.append('>')
    lines.append(f'> 页内的 `{mark}` 为点击分隔符。本文档用于阅读与打印，如需修改后回导请使用 Word 或 JSON 导出。')
    lines.append('')

    lines.append('| 页 | 段数 | 字数 | 时长 |')
    lines.append('| --- | --- | --- | --- |')
    for page, texts, duration, words in rows:
        lines.append(f'| {page} | {len(texts)} | {words} | {format_duration(duration)} |')
    lines.append('')

    for page, texts, duration, _words in rows:
        heading = f'## 第 {page} 页'
        if duration > 0:
            heading += f'（{format_duration(duration)}）'
        lines.append(heading)
        lines.append('')

        body = [text.strip() for text in texts]
        if not any(body):
            lines.append('*（本页无讲稿）*')
            lines.append('')
            continue

        for index, text in enumerate(body):
            if index:
                lines.append(mark)
                lines.append('')
            lines.append(text if text else '*（空白分段）*')
            lines.append('')

    return '\n'.join(lines).rstrip() + '\n'


def build_srt(notes_list, durations):
    """按音频实际时长连续排布并生成 SRT 字幕。"""
    blocks = []
    cursor = 0.0
    index = 0

    for position, note in enumerate(notes_list):
        duration = float(durations[position]) if position < len(durations) else 0.0
        text = str(note.get('text', '')).strip()
        start = cursor
        cursor += duration

        # 空白段仍占用播放时间；只省略字幕块，不能省略时间轴推进。
        if not text or duration <= 0:
            continue

        index += 1
        blocks.append(
            f'{index}\n'
            f'{format_srt_timestamp(start)} --> {format_srt_timestamp(cursor)}\n'
            f'{text}\n'
        )

    return '\n'.join(blocks)


def build_script_json(notes_list, durations, *, source_name='', mark='●',
                      engine_name='', speaker=''):
    """构造可回导的讲稿 JSON 数据，并以分段列表保留页内点击节奏。"""
    pages = []
    for page, texts, duration, words in page_summary(notes_list, durations):
        pages.append({
            'page': page,
            'segments': texts,
            'duration': round(float(duration), 3),
            'words': words,
        })

    return {
        'format': script_io.JSON_FORMAT,
        'version': script_io.JSON_VERSION,
        'source_name': source_name,
        'mark': mark,
        'exported_at': datetime.now().isoformat(timespec='seconds'),
        'engine': engine_name,
        'speaker': speaker,
        'pages': pages,
    }


def export_audio_files(media_paths, notes_list, target_dir: Path) -> int:
    """逐条复制音频并按顺序和页码命名，返回成功导出数。"""
    target_dir.mkdir(parents=True, exist_ok=True)
    exported = 0

    for index, path in enumerate(media_paths):
        source = Path(path)
        if not source.exists():
            print(f'[导出] 音频缺失，已跳过：{source}')
            continue
        page = int(notes_list[index].get('page', 0)) if index < len(notes_list) else 0
        target = target_dir / f'{index + 1:03d}_第{page}页{source.suffix.lower()}'
        shutil.copy2(source, target)
        exported += 1

    return exported


def merge_audio(media_paths, target_path: Path) -> None:
    """在不解码、不转码的前提下合并同格式音频。

    WAV 输入必须具有相同声道数、采样宽度和采样率；MP3 输入按文件字节顺序拼接。
    格式混合时抛出 ``RuntimeError``，因为项目未引入音频解码依赖。
    """
    sources = [Path(p) for p in media_paths if Path(p).exists()]
    if not sources:
        raise RuntimeError('没有可合并的音频')

    suffixes = {path.suffix.lower() for path in sources}
    if len(suffixes) > 1:
        joined = '、'.join(sorted(suffixes))
        raise RuntimeError(f'音频格式不一致（{joined}），无法在不转码的前提下合并，请改用多文件导出。')

    suffix = suffixes.pop()
    target_path.parent.mkdir(parents=True, exist_ok=True)

    if suffix == '.wav':
        params = None
        with wave.open(str(target_path), 'wb') as writer:
            for path in sources:
                with wave.open(str(path), 'rb') as reader:
                    current = reader.getparams()
                    if params is None:
                        params = current
                        writer.setnchannels(current.nchannels)
                        writer.setsampwidth(current.sampwidth)
                        writer.setframerate(current.framerate)
                    elif (current.nchannels, current.sampwidth, current.framerate) != (
                            params.nchannels, params.sampwidth, params.framerate):
                        raise RuntimeError(f'音频参数不一致（{path.name}），请改用多文件导出。')
                    writer.writeframes(reader.readframes(reader.getnframes()))
        return

    with target_path.open('wb') as out:
        for path in sources:
            out.write(path.read_bytes())

def _set_repeating_header(row):
    """把表头标记为跨页重复，长讲稿翻页后仍能看到列含义。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement('w:tblHeader'))


def write_docx(target_path: Path, notes_list, durations, *, mark='●'):
    """将讲稿写为可稳定回导的 Word 表格。

    整篇文档只有一张表：页码和正文由单元格边界承载，因此正文编辑不会破坏
    分页结构；时长列仅供阅读，回导时不参与解析。
    """
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, RGBColor

    document = Document()

    # 收窄页边距，把宽度尽量让给讲稿列
    section = document.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    usable = section.page_width - section.left_margin - section.right_margin

    page_width = Inches(0.5)
    time_width = Inches(0.62)
    body_width = usable - page_width - time_width
    widths = (page_width, body_width, time_width)

    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = False

    header = table.rows[0]
    for index, label in enumerate(('页码', '讲稿', '时长')):
        cell = header.cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _apply_font(paragraph.add_run(label), size=10, bold=True)
    _set_repeating_header(header)

    # 页码与时长用小字号和灰色弱化，让讲稿正文成为视觉重心
    muted = RGBColor(0x60, 0x60, 0x60)

    for page, texts, duration, _words in page_summary(notes_list, durations):
        row = table.add_row().cells

        row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        page_paragraph = row[0].paragraphs[0]
        page_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _apply_font(page_paragraph.add_run(str(page)), size=10, color=muted)

        body = mark.join(text.strip() for text in texts)
        body = _CONTROL_CHARS_RE.sub('', body)
        paragraph = row[1].paragraphs[0]
        for line_index, line in enumerate(body.split('\n')):
            if line_index:
                paragraph = row[1].add_paragraph()
            _apply_font(paragraph.add_run(line), size=11)

        row[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        time_paragraph = row[2].paragraphs[0]
        time_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _apply_font(time_paragraph.add_run(format_duration(duration)), size=9, color=muted)

    # 列宽必须逐单元格设置，只设 columns 的宽度在 Word 中不会生效
    for index, width in enumerate(widths):
        table.columns[index].width = width
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width

    document.save(str(target_path))
